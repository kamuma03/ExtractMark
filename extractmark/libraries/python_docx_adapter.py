"""LIB-06 -- python-docx adapter."""

from __future__ import annotations

import logging
import time
from pathlib import Path

from extractmark.types import PageInput, PageOutput

logger = logging.getLogger(__name__)


class PythonDocxAdapter:
    """Adapter for python-docx (LIB-06) -- direct Word object model access."""

    lib_id = "LIB-06"

    def process_page(self, page: PageInput) -> PageOutput:
        return PageOutput(
            document_id=page.document_id,
            page_number=page.page_number,
            raw_text="",
            metadata={"note": "python-docx works on DOCX files. Use process_document()."},
        )

    def process_document(self, doc_path: Path) -> list[PageOutput]:
        from docx import Document

        document_id = doc_path.stem
        start = time.perf_counter()

        try:
            doc = Document(str(doc_path))
        except Exception as e:
            logger.error("python-docx failed to open %s: %s", doc_path, e)
            return []

        # Extract paragraphs
        text_parts: list[str] = []
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level_num = int(level)
                except ValueError:
                    level_num = 1
                text_parts.append("#" * level_num + " " + para.text)
            elif para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        tables_md: list[str] = []
        for table in doc.tables:
            md_rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                md_rows.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    md_rows.append("| " + " | ".join("---" for _ in cells) + " |")
            if md_rows:
                tables_md.append("\n".join(md_rows))

        elapsed_ms = (time.perf_counter() - start) * 1000
        full_text = "\n\n".join(text_parts)
        if tables_md:
            full_text += "\n\n" + "\n\n".join(tables_md)

        return [PageOutput(
            document_id=document_id,
            page_number=0,
            raw_text=full_text,
            tables=tables_md,
            inference_time_ms=elapsed_ms,
            metadata={"library": self.lib_id, "paragraphs": len(text_parts)},
        )]
